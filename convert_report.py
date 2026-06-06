"""
Convert Report.md sang Report.docx với định dạng chuẩn khóa luận:
- Font: Times New Roman 13pt (body), 14-16pt (heading)
- Heading 1 = Chương (center, bold, uppercase)
- Heading 2 = Mục (bold)
- Heading 3 = Tiểu mục (bold, italic)
- Bảng: grid style
- Đánh số trang
"""
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = Path("docs/Report.md")
DST  = Path("docs/Report.docx")
BAK  = Path(f"docs/backups/Report_backup_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")

# Backup if exists
BAK.parent.mkdir(parents=True, exist_ok=True)
if DST.exists():
    shutil.copy2(DST, BAK)
    print(f"Backup: {BAK}")

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────────────────
styles = doc.styles

def set_font(style, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold   = bold
    style.font.italic = italic
    if color:
        style.font.color.rgb = RGBColor(*color)
    # Also set East Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.append(rFonts)

# Normal text
normal = styles['Normal']
set_font(normal)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = Pt(20)

# Heading 1 — Chương
h1 = styles['Heading 1']
set_font(h1, size=16, bold=True, color=(26, 82, 118))
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after  = Pt(12)

# Heading 2
h2 = styles['Heading 2']
set_font(h2, size=14, bold=True, color=(26, 82, 118))
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after  = Pt(8)

# Heading 3
h3 = styles['Heading 3']
set_font(h3, size=13, bold=True)
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after  = Pt(6)

# ── Read and parse MD ───────────────────────────────────────────────────────
lines = SRC.read_text(encoding='utf-8').splitlines()

def strip_md(text):
    """Remove markdown formatting for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'^\\\-\s', '- ', text)
    text = re.sub(r'^\\\*\s', '* ', text)
    return text.strip()

def add_run_with_bold(para, text):
    """Add text with **bold** inline support."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part:
            para.add_run(strip_md(part).replace('\\-', '-').replace('\\*', '*'))

i = 0
in_table = False
table_rows = []
table_header = None

while i < len(lines):
    raw = lines[i]
    line = raw.strip()

    # ── HEADING ──
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip().lstrip('\\').strip()
        if title:
            p = doc.add_heading(title.upper(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    if line.startswith('## '):
        doc.add_heading(strip_md(line[3:]), level=2)
        i += 1
        continue

    if line.startswith('### '):
        doc.add_heading(strip_md(line[4:]), level=3)
        i += 1
        continue

    if line.startswith('#### '):
        p = doc.add_paragraph()
        r = p.add_run(strip_md(line[5:]))
        r.bold = True
        r.font.size = Pt(13)
        i += 1
        continue

    # ── TABLE ──
    if line.startswith('|'):
        if not in_table:
            in_table = True
            table_rows = []
            table_header = None
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        table_rows.append(cells)
        i += 1
        # Peek if next line is still table
        if i < len(lines) and lines[i].strip().startswith('|'):
            continue
        else:
            # Flush table
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        if ci < max_cols:
                            cell = tbl.cell(ri, ci)
                            p = cell.paragraphs[0]
                            p.clear()
                            add_run_with_bold(p, strip_md(cell_text))
                            p.paragraph_format.space_after = Pt(2)
                            for run in p.runs:
                                run.font.size = Pt(11)
                                run.font.name = 'Times New Roman'
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
            doc.add_paragraph()
            in_table = False
            table_rows = []
        continue

    in_table = False

    # ── BLANK ──
    if not line:
        doc.add_paragraph()
        i += 1
        continue

    # ── MATH ($$) — keep as plain text ──
    if line.startswith('$$'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.replace('$$', '').strip())
        r.font.size = Pt(12)
        r.italic = True
        i += 1
        continue

    # ── LIST ──
    if line.startswith('- ') or line.startswith('\\- ') or re.match(r'^\d+\.\s', line):
        clean = re.sub(r'^\\?-\s', '', line)
        clean = re.sub(r'^\d+\.\s', '', clean)
        p = doc.add_paragraph(style='List Bullet')
        add_run_with_bold(p, clean)
        for run in p.runs:
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
        i += 1
        continue

    # ── FIGURE/IMAGE caption ──
    if line.startswith('Hình ') or line.startswith('Bảng '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(strip_md(line))
        r.italic = True
        r.font.size = Pt(12)
        i += 1
        continue

    # ── ITALIC/BOLD caption (_text_) ──
    if line.startswith('_') and line.endswith('_'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(strip_md(line).strip('_'))
        r.italic = True
        r.font.size = Pt(12)
        i += 1
        continue

    # ── NORMAL PARAGRAPH ──
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    add_run_with_bold(p, line)
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
    i += 1

# ── Save ────────────────────────────────────────────────────────────────────
doc.save(str(DST))
print(f"Saved: {DST}  ({DST.stat().st_size // 1024} KB)")
